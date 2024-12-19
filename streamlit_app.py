import asyncio
import json
import logging
import os
import subprocess
import uuid
from datetime import datetime
from functools import partial
from pathlib import Path
from typing import Dict, List, Optional, Sequence

import nest_asyncio
import pandas as pd
import streamlit as st
import venv
from openai import OpenAI
from autogen_agentchat.agents import (
    AssistantAgent,
    BaseChatAgent,
    CodeExecutorAgent,
    UserProxyAgent,
)
from autogen_agentchat.base import Response
from autogen_agentchat.conditions import HandoffTermination, MaxMessageTermination, TextMentionTermination
from autogen_agentchat.messages import AgentMessage, ChatMessage, TextMessage
from autogen_agentchat.teams import MagenticOneGroupChat, SelectorGroupChat
from autogen_core import CancellationToken
from autogen_core.models import ChatCompletionClient, SystemMessage, UserMessage
from autogen_core.tools import FunctionTool
from autogen_ext.code_executors.local import LocalCommandLineCodeExecutor
from autogen_ext.models.openai import OpenAIChatCompletionClient

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Apply nest_asyncio to allow nested event loops
nest_asyncio.apply()
loop = asyncio.get_event_loop()

# Set OpenAI API key securely using Streamlit's secrets management
os.environ["OPENAI_API_KEY"] = st.secrets["openai"]["api_key"]


# Initialize OpenAI clients
openai_client = OpenAI()
model_client = OpenAIChatCompletionClient(model="gpt-4o")

# Define base directories
BASE_DIRECTORY = Path(__file__).resolve().parent / "working"
PROJECTS_DIR = BASE_DIRECTORY / "projects"

# Ensure the projects directory exists
PROJECTS_DIR.mkdir(parents=True, exist_ok=True)

from typing import Optional, Dict, Any, List
from datetime import datetime
import uuid

class ChatMessage:
    """Represents a chat message with enhanced features."""
    def __init__(
        self,
        role: str,
        content: str,
        message_type: str = "text",
        metadata: Optional[Dict[str, Any]] = None
    ):
        self.id = str(uuid.uuid4())
        self.role = role
        self.content = content
        self.message_type = message_type  # text, code, system, error
        self.metadata = metadata or {}
        self.timestamp = datetime.now()
        self.reactions: List[str] = []

    def to_dict(self) -> Dict:
        return {
            "id": self.id,
            "role": self.role,
            "content": self.content,
            "message_type": self.message_type,
            "metadata": self.metadata,
            "timestamp": self.timestamp.isoformat(),
            "reactions": self.reactions
        }

    @staticmethod
    def from_dict(data: Dict) -> 'ChatMessage':
        msg = ChatMessage(
            role=data["role"],
            content=data["content"],
            message_type=data["message_type"],
            metadata=data["metadata"]
        )
        msg.id = data["id"]
        msg.timestamp = datetime.fromisoformat(data["timestamp"])
        msg.reactions = data["reactions"]
        return msg

def render_message_header(message: ChatMessage):
    """Renders the header of a chat message."""
    col1, col2 = st.columns([6, 1])
    
    with col1:
        role_icons = {
            "user": "👤",
            "assistant": "🤖",
            "system": "⚙️",
            "error": "⚠️"
        }
        icon = role_icons.get(message.role, "📝")
        st.markdown(f"{icon} **{message.role.title()}** · {message.timestamp.strftime('%I:%M %p')}")
    
    with col2:
        if message.role == "assistant":
            st.button("🔄", key=f"regenerate_{message.id}", help="Regenerate response")

def render_code_message(content: str, language: str = "python"):
    """Renders a code message with syntax highlighting."""
    st.code(content, language=language)
    
    col1, col2 = st.columns([1, 6])
    with col1:
        if st.button("📋", key=f"copy_{uuid.uuid4()}", help="Copy to clipboard"):
            st.toast("Code copied to clipboard!", icon="📋")

def render_error_message(content: str):
    """Renders an error message with warning styling."""
    st.error(content)

def render_system_message(content: str):
    """Renders a system message with info styling."""
    st.info(content)

def render_chat_message(message: ChatMessage):
    """Renders a chat message with appropriate styling based on type."""
    with st.chat_message(message.role):
        # Message header with timestamp and actions
        render_message_header(message)
        
        # Message content based on type
        if message.message_type == "code":
            render_code_message(
                message.content,
                message.metadata.get("language", "python")
            )
        elif message.message_type == "error":
            render_error_message(message.content)
        elif message.message_type == "system":
            render_system_message(message.content)
        else:
            st.markdown(message.content)
        
        # Message feedback (for assistant messages)
        if message.role == "assistant":
            col1, col2, col3 = st.columns([1, 1, 5])
            with col1:
                if st.button("👍", key=f"helpful_{message.id}", help="Mark as helpful"):
                    if "👍" not in message.reactions:
                        message.reactions.append("👍")
                        st.toast("Marked as helpful!", icon="👍")
            with col2:
                if st.button("👎", key=f"unhelpful_{message.id}", help="Mark as unhelpful"):
                    if "👎" not in message.reactions:
                        message.reactions.append("👎")
                        st.toast("Feedback recorded", icon="👎")

def render_message_content(message: ChatMessage):
    """Renders message content with enhanced styling."""
    # Message timestamp and actions header
    col1, col2 = st.columns([6, 1])
    with col1:
        st.caption(f"{message.timestamp.strftime('%I:%M %p')}")
    
    # Message content
    if message.message_type == "code":
        st.code(message.content, language=message.metadata.get("language", "python"))
        if st.button("📋 Copy", key=f"copy_{message.id}"):
            st.toast("Code copied to clipboard!", icon="📋")
    elif message.message_type == "error":
        st.error(message.content)
    elif message.message_type == "system":
        st.info(message.content)
    else:
        st.markdown(message.content)
    
    # Message feedback (for assistant messages only)
    if message.role == "assistant":
        col1, col2, col3, col4 = st.columns([1, 1, 1, 4])
        with col1:
            if st.button("👍", key=f"helpful_{message.id}"):
                if "👍" not in message.reactions:
                    message.reactions.append("👍")
                    st.toast("Thank you for your feedback!", icon="👍")
        with col2:
            if st.button("👎", key=f"unhelpful_{message.id}"):
                if "👎" not in message.reactions:
                    message.reactions.append("👎")
                    st.toast("Feedback recorded", icon="👎")
        with col3:
            if st.button("🔄", key=f"regenerate_{message.id}"):
                st.session_state.regenerating = message.id
                st.rerun()

def render_welcome_guidance():
    """Renders the welcome message and user guidance."""
    st.markdown("""
        <div style='text-align: center; margin: 40px 0;'>
            <h1>👋 Welcome to your Research Project</h1>
            <p style='font-size: 1.2em; margin: 20px 0;'>
                I'm your AI research assistant. I can help you create and manage research briefs.
            </p>
        </div>
    """, unsafe_allow_html=True)

    # Expandable sections with guidance
    with st.expander("🎯 What I Can Help You With", expanded=True):
        st.markdown("""
            - **Create Research Briefs**: I'll guide you through creating comprehensive research briefs
            - **Refine Requirements**: I'll help clarify and improve your research objectives
            - **Generate Documentation**: I can create Word documents with proper formatting
            - **Version Management**: Keep track of different versions of your research briefs
            """)

    with st.expander("💡 Example Prompts", expanded=True):
        for prompt in [
            "Help me create a research brief for market analysis",
            "I need a research brief about consumer behavior in e-commerce",
            "Can you help me write a brief for competitive analysis?",
            "I want to study emerging trends in renewable energy",
            "Generate a research brief for social media marketing strategy"
        ]:
            if st.button(prompt, key=f"example_{prompt}", use_container_width=True):
                return prompt

    with st.expander("📝 Research Brief Components", expanded=False):
        st.markdown("""
            A complete research brief typically includes:
            1. **Research Objective**: What you aim to achieve
            2. **Company Background**: Context about your organization
            3. **Business Context**: The problem or opportunity
            4. **Research Questions**: Key questions to answer
            5. **Stakeholders**: Who will use the findings
            6. **Key Decisions**: What decisions this will inform
            7. **Methodology**: Suggested research approaches
            8. **Timeline**: Project milestones
            9. **Deliverables**: Expected outputs
            10. **Budget**: Cost considerations
            
            Don't worry if you don't have all these details - I can help you figure them out!
            """)

    with st.expander("❓ How to Use This Assistant", expanded=False):
        st.markdown("""
            1. **Start a Conversation**: Just type your request or click an example prompt
            2. **Answer Questions**: I'll ask for details to create a complete brief
            3. **Review & Refine**: We'll review the content together
            4. **Generate Document**: Once approved, I'll create a Word document
            5. **Version Control**: Access different versions from the sidebar
            
            📌 **Tip**: The more details you provide, the better I can help!
            """)

def add_chat_guidance_button():
    """Adds a help button that shows guidance when clicked."""
    if st.button("ℹ️ Show Help", key="show_help"):
        st.session_state.show_guidance = True
        st.rerun()

def render_chat_interface(project_id: str):
    """Renders the enhanced chat interface with user guidance."""
    # Display message history
    if not st.session_state.messages[project_id]:
        # Show welcome message and guidance for new chats
        example_prompt = render_welcome_guidance()
        if example_prompt:
            handle_user_input(example_prompt, project_id)
            st.rerun()
    else:
        # Add help button for existing chats
        add_chat_guidance_button()
        
        if st.session_state.get('show_guidance', False):
            render_welcome_guidance()
            if st.button("Hide Help", key="hide_help"):
                st.session_state.show_guidance = False
                st.rerun()
        
        # Display message history
        for message in st.session_state.messages[project_id]:
            if isinstance(message, dict):
                message = ChatMessage(
                    role=message["role"],
                    content=message["content"],
                    message_type="text"
                )
            
            with st.chat_message(message.role, avatar="👤" if message.role == "user" else "🤖"):
                render_message_content(message)

    # Chat input
    user_input = st.chat_input("Message the Research Assistant...")
    if user_input:
        handle_user_input(user_input, project_id)

def handle_user_input(user_input: str, project_id: str):
    """Handles user input and generates assistant response."""
    # Add user message
    user_message = ChatMessage(
        role="user",
        content=user_input,
        message_type="text"
    )
    st.session_state.messages[project_id].append(user_message)
    
    # Process with assistant
    try:
        assistant_response = process_assistant_message(user_input, project_id)
        
        # Add assistant message
        assistant_message = ChatMessage(
            role="assistant",
            content=assistant_response.get("content", ""),
            message_type=assistant_response.get("type", "text"),
            metadata=assistant_response.get("metadata", {})
        )
        st.session_state.messages[project_id].append(assistant_message)
        
    except Exception as e:
        # Add error message
        error_message = ChatMessage(
            role="error",
            content=f"An error occurred: {str(e)}",
            message_type="error"
        )
        st.session_state.messages[project_id].append(error_message)
    
    st.rerun()

def process_assistant_message(user_input: str, project_id: str) -> Dict[str, Any]:
    """
    Processes user input through the assistant agent.
    Returns a dictionary with the response details.
    """
    assistant_agent = st.session_state.assistant_agent
    # Ensure current project ID is set in session state
    st.session_state.current_project_id = project_id
    
    # Show processing indicator
    with st.spinner("Assistant is thinking..."):
        try:
            result = loop.run_until_complete(
                assistant_agent.on_messages(
                    [TextMessage(content=user_input, source="user")],
                    cancellation_token=CancellationToken()
                )
            )
            
            if result.chat_message:
                return {
                    "content": result.chat_message.content,
                    "type": "text",
                    "metadata": {
                        "project_id": project_id
                    }
                }
            else:
                return {
                    "content": "No response generated.",
                    "type": "error",
                    "metadata": {
                        "project_id": project_id
                    }
                }
                
        except Exception as e:
            logger.error(f"Error processing message: {str(e)}")
            return {
                "content": f"Error processing message: {str(e)}",
                "type": "error",
                "metadata": {
                    "error_type": type(e).__name__,
                    "project_id": project_id
                }
            }
        
from datetime import datetime
import json
import uuid
from pathlib import Path
from typing import Dict, List, Optional
import bcrypt

class User:
    """Represents a user in the system."""
    def __init__(self, username: str, email: str, password_hash: str, role: str = "user"):
        self.id = str(uuid.uuid4())
        self.username = username
        self.email = email
        self.password_hash = password_hash
        self.role = role  # "admin" or "user"
        self.created_at = datetime.now()
        self.last_login = None
        self.project_ids: List[str] = []

    def to_dict(self) -> Dict:
        """Serializes the user to a dictionary."""
        return {
            "id": self.id,
            "username": self.username,
            "email": self.email,
            "password_hash": self.password_hash,
            "role": self.role,
            "created_at": self.created_at.isoformat(),
            "last_login": self.last_login.isoformat() if self.last_login else None,
            "project_ids": self.project_ids
        }

    @staticmethod
    def from_dict(data: Dict) -> 'User':
        """Creates a User instance from a dictionary."""
        user = User(
            username=data["username"],
            email=data["email"],
            password_hash=data["password_hash"],
            role=data["role"]
        )
        user.id = data["id"]
        user.created_at = datetime.fromisoformat(data["created_at"])
        user.last_login = datetime.fromisoformat(data["last_login"]) if data["last_login"] else None
        user.project_ids = data["project_ids"]
        return user

class UserManagement:
    """Handles user management operations."""
    def __init__(self, base_dir: Path):
        self.base_dir = base_dir
        self.users_file = base_dir / "users.json"
        self.users: Dict[str, User] = {}
        self.load_users()

    def load_users(self):
        """Loads users from the JSON file."""
        if self.users_file.exists():
            with open(self.users_file, "r") as f:
                data = json.load(f)
                self.users = {
                    uid: User.from_dict(udata)
                    for uid, udata in data.items()
                }

    def save_users(self):
        """Saves users to the JSON file."""
        with open(self.users_file, "w") as f:
            json.dump(
                {uid: user.to_dict() for uid, user in self.users.items()},
                f,
                indent=2
            )

    def create_user(self, username: str, email: str, password: str, role: str = "user") -> Optional[User]:
        """Creates a new user."""
        if any(u.username == username for u in self.users.values()):
            return None
        
        password_hash = bcrypt.hashpw(password.encode(), bcrypt.gensalt()).decode()
        user = User(username, email, password_hash, role)
        self.users[user.id] = user
        self.save_users()
        return user

    def authenticate_user(self, username: str, password: str) -> Optional[User]:
        """Authenticates a user."""
        user = next((u for u in self.users.values() if u.username == username), None)
        if user and bcrypt.checkpw(password.encode(), user.password_hash.encode()):
            user.last_login = datetime.now()
            self.save_users()
            return user
        return None

    def assign_project_to_user(self, user_id: str, project_id: str):
        """Assigns a project to a user."""
        if user_id in self.users:
            if project_id not in self.users[user_id].project_ids:
                self.users[user_id].project_ids.append(project_id)
                self.save_users()

    def remove_project_from_user(self, user_id: str, project_id: str):
        """Removes a project from a user."""
        if user_id in self.users and project_id in self.users[user_id].project_ids:
            self.users[user_id].project_ids.remove(project_id)
            self.save_users()

    def get_user_projects(self, user_id: str) -> List[str]:
        """Gets all project IDs assigned to a user."""
        return self.users[user_id].project_ids if user_id in self.users else []

def render_login_page():
    """Renders the login page."""
    st.title("Login")
    
    username = st.text_input("Username")
    password = st.text_input("Password", type="password")
    
    col1, col2 = st.columns([1, 1])
    with col1:
        if st.button("Login"):
            user = st.session_state.user_manager.authenticate_user(username, password)
            if user:
                st.session_state.current_user = user
                st.session_state.current_user_id = user.id
                st.success(f"Welcome back, {user.username}!")
                st.rerun()
            else:
                st.error("Invalid username or password")
    
    with col2:
        if st.button("Register"):
            st.session_state.show_registration = True
            st.rerun()

def render_registration_page():
    """Renders the registration page."""
    st.title("Register New User")
    
    username = st.text_input("Username")
    email = st.text_input("Email")
    password = st.text_input("Password", type="password")
    confirm_password = st.text_input("Confirm Password", type="password")
    
    col1, col2 = st.columns([1, 1])
    with col1:
        if st.button("Register"):
            if password != confirm_password:
                st.error("Passwords do not match")
            elif not username or not email or not password:
                st.error("All fields are required")
            else:
                user = st.session_state.user_manager.create_user(username, email, password)
                if user:
                    st.success("Registration successful! Please login.")
                    st.session_state.show_registration = False
                    st.rerun()
                else:
                    st.error("Username already exists")
    
    with col2:
        if st.button("Back to Login"):
            st.session_state.show_registration = False
            st.rerun()

def initialize_user_management():
    """Initializes the user management system."""
    if 'user_manager' not in st.session_state:
        st.session_state.user_manager = UserManagement(PROJECTS_DIR)
    
    if 'current_user' not in st.session_state:
        st.session_state.current_user = None
    
    if 'current_user_id' not in st.session_state:
        st.session_state.current_user_id = None
    
    if 'show_registration' not in st.session_state:
        st.session_state.show_registration = False

### Utility Classes ###

class VirtualEnvContext:
    """Context holder for virtual environment executables."""

    def __init__(self, env_exe: Path, bin_path: Path):
        self.env_exe = str(env_exe)
        self.bin_path = str(bin_path)


class Document:
    """Represents a document within a project."""

    def __init__(self, name: str, content: str = "", version: str = "v1"):
        self.name = name
        self.content = content
        self.version = version
        self.created_at = datetime.now()
        self.last_modified = datetime.now()

    def to_dict(self) -> Dict:
        return {
            "name": self.name,
            "content": self.content,
            "version": self.version,
            "created_at": self.created_at.isoformat(),
            "last_modified": self.last_modified.isoformat(),
        }

    @staticmethod
    def from_dict(data: Dict) -> 'Document':
        doc = Document(data["name"], data["content"], data["version"])
        doc.created_at = datetime.fromisoformat(data["created_at"])
        doc.last_modified = datetime.fromisoformat(data["last_modified"])
        return doc


class Project:
    """Represents a project with associated documents and user permissions."""

    def __init__(self, name: str, owner_id: str, description: Optional[str] = None, status: str = "Not Started"):
        self.id = str(uuid.uuid4())
        self.name = name
        self.description = description
        self.status = status
        self.owner_id = owner_id
        self.shared_users: List[str] = []  # List of user IDs with access
        self.created_at = datetime.now()
        self.last_modified = datetime.now()
        self.documents: Dict[str, Document] = {}
        self.document_versions: Dict[str, Dict[str, Document]] = {}

    def add_shared_user(self, user_id: str):
        """Adds a user to the shared users list."""
        if user_id not in self.shared_users:
            self.shared_users.append(user_id)

    def remove_shared_user(self, user_id: str):
        """Removes a user from the shared users list."""
        if user_id in self.shared_users:
            self.shared_users.remove(user_id)

    def can_access(self, user_id: str) -> bool:
        """Checks if a user has access to the project."""
        return user_id == self.owner_id or user_id in self.shared_users

    def to_dict(self) -> Dict:
        """Serializes the project to a dictionary."""
        return {
            "id": self.id,
            "name": self.name,
            "description": self.description,
            "status": self.status,
            "owner_id": self.owner_id,
            "shared_users": self.shared_users,
            "created_at": self.created_at.isoformat(),
            "last_modified": self.last_modified.isoformat(),
            "documents": {name: doc.to_dict() for name, doc in self.documents.items()},
            "document_versions": {
                name: {ver: doc.to_dict() for ver, doc in versions.items()}
                for name, versions in self.document_versions.items()
            },
        }

    @staticmethod
    def from_dict(data: Dict) -> 'Project':
        """Deserializes a project from a dictionary."""
        project = Project(
            data["name"],
            data["owner_id"],
            data.get("description"),
            data.get("status", "Not Started")
        )
        project.id = data["id"]
        project.shared_users = data.get("shared_users", [])
        project.created_at = datetime.fromisoformat(data["created_at"])
        project.last_modified = datetime.fromisoformat(data["last_modified"])

        # Reconstruct documents
        for name, doc_data in data.get("documents", {}).items():
            doc = Document.from_dict(doc_data)
            project.documents[name] = doc

        # Reconstruct document versions
        for name, versions in data.get("document_versions", {}).items():
            project.document_versions[name] = {}
            for ver, doc_data in versions.items():
                doc = Document.from_dict(doc_data)
                project.document_versions[name][ver] = doc

        return project

### Virtual Environment Setup ###

def setup_virtual_env(project: Project) -> VirtualEnvContext:
    """
    Sets up a virtual environment for the given project and installs 'python-docx'.
    
    Args:
        project (Project): The project for which to set up the virtual environment.
    
    Returns:
        VirtualEnvContext: Paths to the Python executable and bin directory.
    """
    venv_dir = PROJECTS_DIR / project.id / ".venv"

    # Determine paths based on OS
    if os.name == 'nt':
        python_executable = venv_dir / "Scripts/python.exe"
        pip_executable = venv_dir / "Scripts/pip.exe"
    else:
        python_executable = venv_dir / "bin/python"
        pip_executable = venv_dir / "bin/pip"

    # Create virtual environment if it doesn't exist
    if not python_executable.exists():
        st.write(f"Creating virtual environment for project '{project.name}'...")
        try:
            venv.create(venv_dir, with_pip=True)
            st.success(f"✅ Virtual environment created at {venv_dir}.")
        except Exception as e:
            st.error(f"❌ Failed to create virtual environment: {e}")
            logger.error(f"Failed to create virtual environment for project '{project.name}': {e}")
            st.stop()
    else:
        st.write(f"✅ Virtual environment already exists for project '{project.name}'.")

    # Verify the Python executable
    if not python_executable.exists():
        st.error(f"❌ Python executable not found in the virtual environment at {python_executable}.")
        logger.error(f"Python executable not found at {python_executable}.")
        st.stop()

    # Install 'python-docx'
    python_docx_version = "0.8.11"
    st.write(f"🛠️ Installing 'python-docx=={python_docx_version}' in the virtual environment for project '{project.name}'...")
    try:
        result = subprocess.run(
            [str(pip_executable), "install", f"python-docx=={python_docx_version}"],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE
        )
        st.success(f"✅ 'python-docx=={python_docx_version}' installed successfully.")
    except subprocess.CalledProcessError as e:
        st.error(f"❌ Failed to install 'python-docx=={python_docx_version}': {e.stderr.decode()}")
        logger.error(f"Failed to install 'python-docx=={python_docx_version}' for project '{project.name}': {e.stderr.decode()}")
        st.stop()
    except Exception as e:
        st.error(f"❌ Exception occurred while installing 'python-docx=={python_docx_version}': {e}")
        logger.error(f"Exception during 'python-docx' installation for project '{project.name}': {e}")
        st.stop()

    return VirtualEnvContext(env_exe=python_executable, bin_path=python_executable.parent)

### AI Agents and Tools ###

async def call_research_brief_team_tool(summary: str) -> str:
    """
    Calls the research brief writing team with a summary of requirements.
    """
    try:
        # Get the current project ID from session state
        project_id = st.session_state.current_project_id
        if not project_id:
            return "Project ID not found. Please select a project first."

        project = st.session_state.projects.get(project_id)
        if not project:
            return "Project not found. Please ensure the project exists."

        team = create_research_brief_team(model_client, project_id)
        if not team:
            return "Failed to create the research brief team."

        # Create output directory if it doesn't exist
        output_dir = PROJECTS_DIR / project_id / "output"
        output_dir.mkdir(parents=True, exist_ok=True)

        cancellation_token = CancellationToken()
        fixed_instructions = """
            Complete a market research brief that contains the details summarized below. 
            Improve content or phrasing where possible. Omit content that has been flagged as omitted. 
            Always finish with the code_executor to save the finished document in the output directory."""
            
        result = await team.run(task=f"{fixed_instructions}\n\n{summary}", cancellation_token=cancellation_token)

        # Process team response
        team_response = []
        for msg in result.messages:
            if msg.source != "UserProxy":
                team_response.append(f"**{msg.source}**: {msg.content}")

        return "\n\n".join(team_response)

    except Exception as e:
        logger.error(f"Error in call_research_brief_team_tool: {e}")
        return f"An error occurred while processing the research brief: {str(e)}"


call_research_brief_tool = FunctionTool(
    call_research_brief_team_tool,
    description="Call the research brief writing team with a summary of requirements."
)


class FrontlineAssistant(AssistantAgent):
    """
    First point of contact agent that gathers user requirements before creating a research brief.
    """

    def __init__(self, model_client: ChatCompletionClient):
        super().__init__(
            name="AI_Assistant",
            description="A helpful assistant that gathers user requirements before creating a research brief",
            tools=[call_research_brief_tool],
            model_client=model_client,
            system_message="""
            You are a helpful AI assistant that converses with a user to help them complete a market research project.

            You gather requirements, explain project findings and progress, and interface with specialized teams of AI agents to complete tasks.
            
            Stay focused on the research project and do not deviate from this task. 

            The research project consists of the following:
                - Research Brief: a document that outlines project requirements

            *Research Brief Instructions*
                When the user asks about generating a research brief your job is to gather the document requirements and then send them to the Research Brief Team.
                Ask about research brief requirements. If the User does not know a requirement ask if they would like the AI Research Brief Team to make a recommendation or to leave it blank. 
                Once all requirements have been reviewed, summarize the results and ask the user to confirm if the document should be produced. Warn the user that it could take a few minutes. Make changes in response to feedback.
                If the user confirms the content is ready, call on the 'call_research_brief_team' tool to pass a summary of the research brief requirements and user instructions to the Research Brief Team.

            *Research Brief Requirements*
                -Research Objective: Clear and concise statement of what the research aims to achieve.
                -Company Background: Brief introduction to the company and relevant context.
                -Business Context: Explanation of the business problem or opportunity prompting the research.
                -Research Questions: Key questions the research must answer.
                -Stakeholders: Primary users of the research findings.
                -Key Decisions: Critical decisions the research will inform.
                -Preferred Methodologies: Suggested approaches (e.g., qualitative, quantitative, mixed methods).
                -Timelines: Key project milestones and deadlines.
                -Deliverables: Expected outputs (e.g., reports, presentations).
                -Budget: Estimated budget or cost considerations.

            *Research Brief Tools*
                -call_research_brief_team
            """
        )


class StreamlitUserProxyAgent(UserProxyAgent):
    """Handles user interactions via the Streamlit chat interface."""

    def __init__(self, project_id: str):
        super().__init__(
            name="UserProxy",
            description="The client who should be consulted for document requirements and final approval of content.",
            input_func=self._get_user_input
        )
        self.project_id = project_id
        self._message_queue = asyncio.Queue()

    @property
    def produced_message_types(self) -> List[type[ChatMessage]]:
        return [TextMessage]

    async def _get_user_input(self, prompt: str, cancellation_token: Optional[CancellationToken] = None) -> str:
        """Handles fetching user input from the Streamlit interface."""
        try:
            if prompt and prompt.strip():
                with st.chat_message("assistant"):
                    st.markdown(prompt)
                st.session_state.messages[self.project_id].append({
                    "role": "assistant",
                    "content": prompt
                })

            while not (cancellation_token and cancellation_token.cancelled):
                try:
                    user_input = await asyncio.wait_for(
                        self._message_queue.get(),
                        timeout=0.1
                    )
                    return user_input
                except asyncio.TimeoutError:
                    continue

            return "CANCELLED"

        except Exception as e:
            logger.error(f"Error getting user input: {e}")
            return "ERROR: Unable to get user input"

    def add_user_message_sync(self, message: str):
        """Adds a user message to the queue synchronously."""
        asyncio.run_coroutine_threadsafe(self._message_queue.put(message), loop)

    async def on_reset(self, cancellation_token: Optional[CancellationToken] = None) -> None:
        """Resets the agent's message queue."""
        self._message_queue = asyncio.Queue()

    async def on_messages(
        self,
        messages: Sequence[AgentMessage],
        cancellation_token: Optional[CancellationToken] = None
    ) -> Response:
        """Handles incoming messages."""
        if messages:
            last_message = messages[-1]
            if last_message.content and last_message.content.strip():
                prompt = last_message.content
                response = await self._get_user_input(prompt, cancellation_token)
                return Response(chat_message=TextMessage(content=response, source=self.name))

        # Await user input if no messages are present
        response = await self._get_user_input("", cancellation_token)
        return Response(chat_message=TextMessage(content=response, source=self.name))


def create_research_brief_team(model_client: ChatCompletionClient, project_id: str) -> Optional[MagenticOneGroupChat]:
    """
    Creates a research brief team using SelectorGroupChat.

    Args:
        model_client (ChatCompletionClient): The AI model client.
        project_id (str): The ID of the current project.

    Returns:
        Optional[MagenticOneGroupChat]: The created team or None if project not found.
    """
    project = st.session_state.projects.get(project_id)
    if not project:
        st.error(f"Project with ID '{project_id}' not found.")
        return None

    Research_Brief_Expert = AssistantAgent(
        name="Research_Brief_Expert",
        description="An expert at writing research briefs that reviews and assesses its contents",
        model_client=model_client,
        system_message="""
        Role & Responsibility:
        You are the Research Brief Expert. Your role is to improve and clarify the components of a market research brief. 
    
        For each of the Required Contents in a market research brief provided to you, you assess the following:
        -Needs to be rephrased for clarity and professionalism
        -Is insufficient and needs additional detail
        -Is missing and needs to be inferred based on the available detail and best practice
        -Is missing but can be omitted based on instructions received

        Provide a restatement of the content and your assessment of it.
        
        Here is the Required Content:

        -Research Objective: Clear and concise statement of what the research aims to achieve.
        -Company Background: Brief introduction to the company and relevant context.
        -Business Context: Explanation of the business problem or opportunity prompting the research.
        -Research Questions: Key questions the research must answer.
        -Stakeholders: Primary users of the research findings.
        -Key Decisions: Critical decisions the research will inform.
        -Preferred Methodologies: Suggested approaches (e.g., qualitative, quantitative, mixed methods).
        -Timelines: Key project milestones and deadlines.
        -Deliverables: Expected outputs (e.g., reports, presentations).
        -Budget: Estimated budget or cost considerations.

    """,
    )

    Market_Research_Expert = AssistantAgent(
        name="Market_Research_Expert",
        description="An expert in market research methods that can be consulted on methodology, best practice, and feasibility",
        model_client=model_client,
        system_message="""
        Role & Responsibility:
        You are the Market Research Expert. You provide constructive feedback, advice, and expert knowedge to improve the quality of Research Briefs.
        When you revieve the content of a research brief with comments you use those comments to improve the content. 
        You never change the meaning of the content but you may improve wording, add missing details, or recommend alternatives.

    """,
    )

    Brief_Writer = AssistantAgent(
        name="Brief_Writer",
        description="A agent that knows how to write market research project briefs",
        model_client=model_client,
        system_message="""
        Role & Responsibility:
        You are the Brief Writer. You synthesize inputs into a complete research brief using the required structure. 
        You only use the content provided to you. It is ok to leave sections of the report blank if the content is missing.

        Example Brief Format:

        ### Research Brief: Smart Lighting Product Launch
        #### 1. Research Objective  
        Understand customer perceptions of the new smart lighting product line.

        #### 2. Company Background  
        XYZ Corp is a leader in smart home devices...

        #### 3. Business Context  
        The company is preparing to launch a new line of smart lighting products...

        #### 4. Research Questions  
        1. How do customers perceive the new product features?  
        2. What barriers exist for adoption?

        #### 5. Stakeholders  
        - Marketing Team  
        - Product Development Team

        #### 6. Key Decisions  
        - Marketing strategy adjustments  
        - Feature prioritization

        #### 7. Preferred Methodology
        - Market Landscape
        - Focus Groups

        #### 8. Timelines
        - 8-12 weeks

        #### 9. Timelines
        - Report with recommendations
        - Presentation

        #### 10. Budget
        -  $30,000 - $40,000
        """,
        )

    Brief_Formatter = AssistantAgent(
        name="Brief_Formatter",
        description="Annotates Briefs for conversion into Word documents",
        model_client=model_client,
        system_message="""
        Role & Responsibility:
        You are the Brief Formatter. You annotate the brief for conversion to a Word document.

        Example Annotations:

        {{H1}} Research Brief: Smart Lighting Product Launch

        {{H2}} Research Objective  
        {{P}} Understand customer perceptions of the new smart lighting product line.

        {{H2}} Company Background  
        {{P}} XYZ Corp is a leader in smart home devices...

        {{BULLET}} How do customers perceive the new product features?  
        {{BULLET}} What barriers exist for adoption?
        """,
        )

    Code_Generator = AssistantAgent(
        name="Code_Generator",
        description="Generates Python code using python-docx to convert annotated briefs into Word documents",
        model_client=model_client,
        system_message="""
        Role & Responsibility:
        You generate Python code using python-docx to convert annotated briefs into Word documents.

        Save the generated document in the 'output' folder inside the project directory.

        Example Conversion to Python Code:

        from docx import Document
        import os

        def create_brief():
            doc = Document()

            doc.add_heading("Research Brief: Smart Lighting Product Launch", level=1)

            doc.add_heading("Research Objective", level=2)
            doc.add_paragraph("Understand customer perceptions of the new smart lighting product line.")

            doc.add_heading("Company Background", level=2)
            doc.add_paragraph("XYZ Corp is a leader in smart home devices...")

            doc.add_heading("Research Questions", level=2)
            questions = ["How do customers perceive the new product features?", "What barriers exist for adoption?"]
            for q in questions:
                doc.add_paragraph(q, style='List Bullet')

            # Ensure the output directory exists
            output_dir = os.path.join(os.getcwd(), 'output')
            os.makedirs(output_dir, exist_ok=True)

            # Save the document in the output directory
            output_path = os.path.join(output_dir, "research_brief.docx")
            doc.save(output_path)
            print(f"Document saved at: {output_path}")

        create_brief()
        """,
    )
    # Setup virtual environment for the project
    venv_context = setup_virtual_env(project)

    # Initialize the local code executor with the virtual environment context
    local_executor = LocalCommandLineCodeExecutor(
        work_dir=str(Path(PROJECTS_DIR) / project_id),
        virtual_env_context=venv_context
    )

    # Initialize the code executor agent
    code_executor_agent = CodeExecutorAgent("code_executor", code_executor=local_executor)

    team = MagenticOneGroupChat(
        [Research_Brief_Expert, Market_Research_Expert, Brief_Writer, Brief_Formatter, Code_Generator, code_executor_agent],
        model_client=model_client)
    
    return team

def get_research_brief_team(project_id: str) -> Optional[MagenticOneGroupChat]:
    """Creates and returns a new research brief team instance."""
    return create_research_brief_team(model_client, project_id)


### User Input Processing ###

async def process_user_input(assistant_agent: FrontlineAssistant, user_input: str, project_id: str):
    """
    Processes user input through the assistant and updates the chat interface.
    
    Args:
        assistant_agent (FrontlineAssistant): The assistant agent.
        user_input (str): The user's input message.
        project_id (str): The current project's ID.
    """
    # Display user message
    with st.chat_message("user"):
        st.markdown(user_input)
    st.session_state.messages[project_id].append({
        "role": "user",
        "content": user_input
    })

    # Get response from the assistant
    assistant_result = await assistant_agent.on_messages(
        [TextMessage(content=user_input, source="user")],
        cancellation_token=CancellationToken()
    )

    # Display the assistant's response
    if assistant_result.chat_message:
        with st.chat_message("assistant"):
            st.markdown(assistant_result.chat_message.content)
        st.session_state.messages[project_id].append({
            "role": "assistant",
            "content": assistant_result.chat_message.content
        })


### Project Management Functions ###

def save_projects():
    """Saves all projects to a JSON file."""
    projects_file = PROJECTS_DIR / "projects.json"
    projects_data = {
        pid: project.to_dict()
        for pid, project in st.session_state.projects.items()
    }
    with open(projects_file, "w") as f:
        json.dump(projects_data, f, indent=2)
    logger.info(f"Saved {len(projects_data)} projects to {projects_file}.")


def load_projects() -> Dict[str, Project]:
    """Loads projects from a JSON file."""
    projects_file = PROJECTS_DIR / "projects.json"
    logger.info(f"Attempting to load projects from {projects_file.resolve()}")
    try:
        with open(projects_file, "r") as f:
            data = json.load(f)
            projects = {
                pid: Project.from_dict(pdata)
                for pid, pdata in data.items()
            }
            logger.info(f"Loaded {len(projects)} projects from {projects_file}.")
            return projects
    except FileNotFoundError:
        logger.info(f"No existing projects found at {projects_file}. Starting fresh.")
        return {}
    except json.JSONDecodeError as e:
        logger.error(f"Error decoding JSON from {projects_file}: {e}. Starting fresh.")
        return {}
    except KeyError as e:
        logger.error(f"Missing key in project data: {e}. Starting fresh.")
        return {}
    except Exception as e:
        logger.error(f"Unexpected error loading projects: {e}. Starting fresh.")
        return {}


### Streamlit Rendering Functions ###

# Suggested prompts and project statuses
SUGGESTED_PROMPTS = [
    "Generate a research brief about market trends in renewable energy",
    "Create a research brief analyzing consumer behavior in e-commerce",
    "Write a research brief on emerging technologies in healthcare",
    "Develop a research brief exploring social media marketing strategies",
    "Draft a research brief investigating sustainability practices in retail"
]

PROJECT_STATUSES = [
    "Not Started",
    "In Progress",
    "Under Review",
    "Completed",
    "On Hold"
]


def render_prompt_pills() -> Optional[str]:
    """Renders suggested prompts as buttons."""
    st.markdown("#### Suggested Prompts")
    cols = st.columns(3)
    for idx, prompt in enumerate(SUGGESTED_PROMPTS):
        col = cols[idx % 3]
        with col:
            if st.button(prompt, key=f"prompt_{idx}", use_container_width=True, type="secondary"):
                return prompt
    return None


def render_document_generator(project_id: str):
    """
    Renders the document generator interface with ChatGPT-style layout.
    """
    # Clear any stale project IDs
    if project_id != st.session_state.get('current_project_id'):
        st.session_state.current_project_id = project_id

    project = st.session_state.projects.get(project_id)
    if not project:
        st.error(f"Project with ID '{project_id}' not found.")
        return

    # Ensure project directories exist
    project_dir = PROJECTS_DIR / project_id
    output_dir = project_dir / "output"
    project_dir.mkdir(parents=True, exist_ok=True)
    output_dir.mkdir(parents=True, exist_ok=True)

    # Initialize states
    st.session_state.messages.setdefault(project_id, [])
    if "assistant_agent" not in st.session_state:
        st.session_state.assistant_agent = FrontlineAssistant(model_client)

    # Sidebar - Document Management
    with st.sidebar:
        # Project Header
        st.title(f"📁 {project.name}")
        status_color = {
            "Not Started": "🔴",
            "In Progress": "🟡",
            "Under Review": "🟣",
            "Completed": "🟢",
            "On Hold": "⚪"
        }.get(project.status, "⚪")
        st.caption(f"{status_color} {project.status}")
        
        if project.owner_id == st.session_state.current_user_id:
            st.caption("👑 Owner")
        else:
            st.caption("👥 Shared")

        st.divider()

        # Generated Documents Section
        st.subheader("📄 Generated Documents")
        output_dir = PROJECTS_DIR / project_id / "output"
        docx_files = list(output_dir.glob("*.docx"))
        
        if docx_files:
            for docx_file in docx_files:
                with st.container():
                    st.markdown(f"📎 **{docx_file.name}**")
                    st.caption(f"Generated: {datetime.fromtimestamp(docx_file.stat().st_mtime).strftime('%Y-%m-%d %H:%M')}")
                    with open(docx_file, "rb") as file:
                        st.download_button(
                            label="📥 Download",
                            data=file,
                            file_name=docx_file.name,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"download_word_{project_id}_{docx_file.stem}"
                        )
                    st.divider()
        else:
            st.info("No documents generated yet")

        # Version History Section
        if project.document_versions.get("Research Brief"):
            st.subheader("📚 Version History")
            versions = project.get_document_versions("Research Brief")
            for version_name, doc in versions.items():
                with st.container():
                    current_marker = " (Current)" if project.documents.get("Research Brief") and project.documents["Research Brief"].version == version_name else ""
                    st.markdown(f"**v{version_name}**{current_marker}")
                    st.caption(f"Modified: {doc.last_modified.strftime('%Y-%m-%d %H:%M')}")
                    
                    col1, col2 = st.columns(2)
                    with col1:
                        if st.button("📌 Set Current", key=f"set_current_{project_id}_{version_name}", type="secondary"):
                            project.documents["Research Brief"] = doc
                            save_projects()
                            st.success(f"Set v{version_name} as current")
                            st.rerun()
                    
                    with col2:
                        if project.owner_id == st.session_state.current_user_id:
                            if st.button("🗑️ Delete", key=f"delete_{project_id}_{version_name}", type="secondary"):
                                del project.document_versions["Research Brief"][version_name]
                                if project.documents.get("Research Brief") and project.documents["Research Brief"].version == version_name:
                                    project.documents.pop("Research Brief", None)
                                save_projects()
                                st.success(f"Deleted v{version_name}")
                                st.rerun()
                    st.divider()

        # Project Info Section (Collapsible)
        with st.expander("ℹ️ Project Information"):
            st.markdown("**Description:**")
            st.markdown(project.description or "*No description provided*")
            st.markdown("**Details:**")
            st.markdown(f"- Created: {project.created_at.strftime('%Y-%m-%d')}")
            st.markdown(f"- Last Modified: {project.last_modified.strftime('%Y-%m-%d %H:%M')}")
            if project.shared_users:
                st.markdown("**Shared with:**")
                for user_id in project.shared_users:
                    user = st.session_state.user_manager.users.get(user_id)
                    if user:
                        st.markdown(f"- {user.username}")

    # Main Chat Area - Full Width
    main_container = st.container()
    with main_container:
        # Custom CSS for chat container
        st.markdown("""
            <style>
                /* Chat container styling */
                .stChatFloatingInputContainer {
                    bottom: 20px;
                    padding: 20px;
                    background: linear-gradient(180deg, rgba(255,255,255,0) 0%, rgba(255,255,255,1) 50%);
                }
                .stChatInput {
                    max-width: 800px;
                    margin: 0 auto;
                }
                /* Message styling */
                .stChatMessage {
                    max-width: 800px;
                    margin: 0 auto;
                    padding: 15px;
                }
                .user-message {
                    background-color: #f7f7f7;
                }
                .assistant-message {
                    background-color: #ffffff;
                }
            </style>
        """, unsafe_allow_html=True)

        # Render chat messages
        messages_container = st.container()
        with messages_container:
            render_chat_interface(project_id)

def render_project_management():
    """Renders the project management interface."""
    st.title("Project Management")
    
    # User info and logout
    col1, col2 = st.columns([3, 1])
    with col2:
        st.markdown(f"Logged in as: **{st.session_state.current_user.username}**")
        if st.button("Logout"):
            st.session_state.current_user = None
            st.session_state.current_user_id = None
            st.rerun()

            
    # Create New Project
    with st.expander("Create New Project", expanded=False):
        col1, col2 = st.columns([2, 1])
        with col1:
            new_project_name = st.text_input("Project Name")
            new_project_desc = st.text_area("Project Description")
        with col2:
            new_project_status = st.selectbox("Initial Status", PROJECT_STATUSES)

        if st.button("Create Project"):
            if new_project_name:
                project = Project(
                    name=new_project_name,
                    owner_id=st.session_state.current_user_id,
                    description=new_project_desc,
                    status=new_project_status
                )
                st.session_state.projects[project.id] = project
                st.session_state.user_manager.assign_project_to_user(
                    st.session_state.current_user_id,
                    project.id
                )
                save_projects()
                st.success(f"Created project: **{new_project_name}**")
                setup_virtual_env(project)
            else:
                st.error("Project name is required")

    # Project List
    st.markdown("### Current Projects")

    # Project filters
    col1, col2 = st.columns([1, 2])
    with col1:
        status_filter = st.multiselect(
            "Filter by Status",
            options=PROJECT_STATUSES,
            default=PROJECT_STATUSES
        )

    # Display projects
    if st.session_state.projects:
        user_projects = st.session_state.user_manager.get_user_projects(st.session_state.current_user_id)
        accessible_projects = [
            project for project in st.session_state.projects.values()
            if project.id in user_projects and project.status in status_filter
        ]

        if accessible_projects:
            for project in accessible_projects:
                with st.container():
                    cols = st.columns([2, 1, 1, 1, 1, 1])

                    with cols[0]:
                        st.markdown(f"**{project.name}**")
                    with cols[1]:
                        st.markdown(f"Status: **{project.status}**")
                    with cols[2]:
                        st.markdown(f"Documents: **{len(project.documents)}**")
                    with cols[3]:
                        if project.owner_id == st.session_state.current_user_id:
                            st.markdown("👑 Owner")
                        else:
                            st.markdown("👥 Shared")
                    with cols[4]:
                        if st.button("Open", key=f"open_{project.id}"):
                            if project.id not in st.session_state.active_tabs:
                                st.session_state.active_tabs.append(project.id)
                            st.session_state.current_tab = project.id
                            st.rerun()
                    with cols[5]:
                        if project.owner_id == st.session_state.current_user_id:
                            if st.button("Share", key=f"share_{project.id}"):
                                st.session_state.sharing_project = project.id

                    # Share Project Section
                    if st.session_state.get('sharing_project') == project.id:
                        with st.expander("Share Project", expanded=True):
                            # Get list of users excluding current user
                            other_users = [
                                user for user in st.session_state.user_manager.users.values()
                                if user.id != st.session_state.current_user_id
                            ]
                            
                            # Display current shared users
                            st.markdown("#### Currently Shared With:")
                            for user_id in project.shared_users:
                                user = st.session_state.user_manager.users.get(user_id)
                                if user:
                                    cols = st.columns([3, 1])
                                    with cols[0]:
                                        st.markdown(f"- {user.username} ({user.email})")
                                    with cols[1]:
                                        if st.button("Remove", key=f"remove_{project.id}_{user.id}"):
                                            project.remove_shared_user(user.id)
                                            st.session_state.user_manager.remove_project_from_user(user.id, project.id)
                                            save_projects()
                                            st.rerun()
                            
                            # Add new shared users
                            st.markdown("#### Share with New Users:")
                            available_users = [
                                user for user in other_users 
                                if user.id not in project.shared_users
                            ]
                            
                            if available_users:
                                selected_user = st.selectbox(
                                    "Select User",
                                    options=available_users,
                                    format_func=lambda x: f"{x.username} ({x.email})"
                                )
                                
                                if st.button("Add User"):
                                    project.add_shared_user(selected_user.id)
                                    st.session_state.user_manager.assign_project_to_user(
                                        selected_user.id,
                                        project.id
                                    )
                                    save_projects()
                                    st.rerun()
                            else:
                                st.info("No more users available to share with")

                            if st.button("Done Sharing"):
                                st.session_state.sharing_project = None
                                st.rerun()
        else:
            st.info("No projects match the selected filters")
    else:
        st.info("No projects created yet")


### Main App ###

def main():
    """Main function to run the Streamlit app."""
    st.set_page_config(layout="wide", page_title="Research Project Manager")

    # Initialize user management
    initialize_user_management()

    # Initialize session state
    if 'active_tabs' not in st.session_state:
        st.session_state.active_tabs = []
    if 'current_tab' not in st.session_state:
        st.session_state.current_tab = None
    if 'messages' not in st.session_state:
        st.session_state.messages = {}
    if 'projects' not in st.session_state:
        st.session_state.projects = load_projects()
    if 'sharing_project' not in st.session_state:
        st.session_state.sharing_project = None
    
    # Handle authentication
    if not st.session_state.current_user:
        if st.session_state.show_registration:
            render_registration_page()
        else:
            render_login_page()
        return

    # Tab Management for authenticated users
    if st.session_state.active_tabs:
        tab_names = ["Project Manager"] + [
            st.session_state.projects[tab_id].name
            for tab_id in st.session_state.active_tabs
            if tab_id in st.session_state.projects
        ]
        tabs = st.tabs(tab_names)

        # Project Manager Tab
        with tabs[0]:
            render_project_management()

        # Project Tabs
        for i, project_id in enumerate(st.session_state.active_tabs, 1):
            if project_id in st.session_state.projects:
                project = st.session_state.projects[project_id]
                if project.can_access(st.session_state.current_user_id):
                    with tabs[i]:
                        col1, col2 = st.columns([6, 1])
                        with col2:
                            if st.button("Close Tab", key=f"close_{project_id}"):
                                st.session_state.active_tabs.remove(project_id)
                                if st.session_state.current_tab == project_id:
                                    st.session_state.current_tab = None
                                st.rerun()
                        render_document_generator(project_id)
    else:
        render_project_management()

if __name__ == "__main__":
    main()